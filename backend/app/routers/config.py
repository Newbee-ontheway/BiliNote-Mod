from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from typing import Optional
import os
from app.utils.response import ResponseWrapper as R

from app.services.cookie_manager import CookieConfigManager
from ffmpeg_helper import ensure_ffmpeg_or_raise

router = APIRouter()
cookie_manager = CookieConfigManager()


class CookieUpdateRequest(BaseModel):
    platform: str
    cookie: str


@router.get("/get_downloader_cookie/{platform}")
def get_cookie(platform: str):
    cookie = cookie_manager.get(platform)
    if not cookie:
        return R.success(msg='未找到Cookies')
    return R.success(
        data={"platform": platform, "cookie": cookie}
    )


@router.post("/update_downloader_cookie")
def update_cookie(data: CookieUpdateRequest):
    cookie_manager.set(data.platform, data.cookie)
    return R.success(

    )

@router.get("/sys_health")
async def sys_health():
    try:
        ensure_ffmpeg_or_raise()
        return R.success()
    except EnvironmentError:
        return R.error(msg="系统未安装 ffmpeg 请先进行安装")

@router.get("/sys_check")
async def sys_check():
    return R.success()


# ==================== 下载地址配置 ====================

import json
from pathlib import Path

_SETTINGS_FILE = Path("config/app_settings.json")


def _load_settings() -> dict:
    if _SETTINGS_FILE.exists():
        return json.loads(_SETTINGS_FILE.read_text(encoding="utf-8"))
    return {}


def _save_settings(settings: dict):
    _SETTINGS_FILE.parent.mkdir(parents=True, exist_ok=True)
    _SETTINGS_FILE.write_text(json.dumps(settings, ensure_ascii=False, indent=2), encoding="utf-8")


class OutputDirRequest(BaseModel):
    output_dir: str


@router.get("/get_output_dir")
def get_output_dir():
    settings = _load_settings()
    current = settings.get("output_dir", os.getenv("NOTE_OUTPUT_DIR", "note_results"))
    return R.success(data={"output_dir": current})


@router.post("/update_output_dir")
def update_output_dir(data: OutputDirRequest):
    path = data.output_dir.strip()
    if not path:
        return R.error(msg="路径不能为空")
    # 尝试创建目录
    try:
        os.makedirs(path, exist_ok=True)
    except Exception as e:
        return R.error(msg=f"无法创建目录: {e}")
    # 持久化
    settings = _load_settings()
    settings["output_dir"] = path
    _save_settings(settings)
    # 运行时更新
    os.environ["NOTE_OUTPUT_DIR"] = path
    return R.success(msg="保存成功")


@router.get("/pick_folder")
def pick_folder():
    """打开 Windows 原生文件夹选择对话框，返回所选路径"""
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)
        folder = filedialog.askdirectory(title="选择下载/导出文件夹")
        root.destroy()
        if folder:
            # 转为 Windows 风格路径
            folder = os.path.normpath(folder)
            return R.success(data={"path": folder})
        return R.success(data={"path": ""})
    except Exception as e:
        return R.error(msg=f"无法打开文件夹选择器: {e}")


# ==================== 导出文件到指定路径 ====================

class ExportFileRequest(BaseModel):
    content: str           # 文本内容或 base64 编码的二进制内容
    filename: str          # 文件名（含扩展名）
    is_base64: bool = False  # True 表示 content 是 base64 编码
    format: str = ""       # 可选：md / pdf / docx — 如需格式转换


def _ensure_pkg(import_name: str, pip_name: str | None = None):
    """检查并自动安装缺失的 Python 包（首次使用时一次性安装）"""
    import importlib
    try:
        importlib.import_module(import_name)
    except ImportError:
        import subprocess, sys
        pkg = pip_name or import_name
        # 先尝试默认源，失败后回退清华镜像
        for extra_args in [[], ["-i", "https://pypi.tuna.tsinghua.edu.cn/simple", "--trusted-host", "pypi.tuna.tsinghua.edu.cn"]]:
            try:
                subprocess.check_call(
                    [sys.executable, "-m", "pip", "install", pkg, "-q", "--disable-pip-version-check"] + extra_args,
                    timeout=120,
                )
                return
            except Exception:
                continue


def _convert_md_to_pdf(md_text: str, out_path) -> bool:
    """将 Markdown 转换为 PDF（使用 markdown_pdf，已内置支持中文）"""
    try:
        _ensure_pkg("markdown_pdf")
        from markdown_pdf import MarkdownPdf, Section

        css = """
body   { font-family: "Microsoft YaHei", "SimHei", sans-serif; font-size: 11pt; line-height: 1.7; color: #222; }
h1     { font-size: 20pt; font-weight: bold; margin-top: 18pt; margin-bottom: 8pt; color: #111; }
h2     { font-size: 16pt; font-weight: bold; margin-top: 14pt; margin-bottom: 6pt; color: #222; }
h3     { font-size: 13pt; font-weight: bold; margin-top: 10pt; margin-bottom: 4pt; }
strong, b { font-weight: bold; }
em, i  { font-style: italic; }
code   { font-family: "Consolas", "Courier New", monospace; font-size: 9pt; background-color: #f0f0f0; }
pre    { font-family: "Consolas", "Courier New", monospace; font-size: 9pt; background-color: #f5f5f5; padding: 8pt; }
blockquote { border-left: 3pt solid #ccc; padding-left: 10pt; color: #555; margin-left: 0; }
table  { border-collapse: collapse; width: 100%; }
th, td { border: 1pt solid #ddd; padding: 5pt 8pt; }
th     { background-color: #f0f0f0; font-weight: bold; }
li     { margin-bottom: 2pt; }
"""
        # 预处理：移除内部锚点链接，避免 pymupdf 处理中文链接时崩溃
        import re as _re
        # [文字](#锚点) → 文字  （保留外部链接）
        cleaned_text = _re.sub(r'\[([^\]]+)\]\(#[^)]*\)', r'\1', md_text)

        pdf = MarkdownPdf(toc_level=0)
        pdf.meta["title"] = ""
        pdf.add_section(Section(cleaned_text, toc=False), user_css=css)
        pdf.save(str(out_path))
        return True
    except Exception as e:
        import traceback
        traceback.print_exc()
        return False


def _convert_md_to_docx(md_text: str, out_path) -> bool:
    """将 Markdown 转换为 Word 文档（基于 markdown-it-py AST，正确处理层级）"""
    try:
        _ensure_pkg("docx", "python-docx")
        from markdown_it import MarkdownIt
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # ── 设置默认字体 ──
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(11)
        normal_style.element.rPr.rFonts.set(qn("w:eastAsia"), "Source Han Serif CN")
        for lvl in range(1, 4):
            hs = doc.styles[f"Heading {lvl}"]
            hs.font.name = "Times New Roman"
            hs.element.rPr.rFonts.set(qn("w:eastAsia"), "Source Han Serif CN")

        # ── 工具函数 ──
        def _set_run_font(run, is_code=False):
            """给 run 设置中英字体"""
            if is_code:
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x88, 0x33, 0x33)
            else:
                run.font.name = "Times New Roman"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "Source Han Serif CN")

        def _add_inline_children(paragraph, children):
            """将 inline token 的 children 渲染为 Word runs"""
            bold = False
            italic = False
            for child in children:
                if child.type == "strong_open":
                    bold = True
                elif child.type == "strong_close":
                    bold = False
                elif child.type == "em_open":
                    italic = True
                elif child.type == "em_close":
                    italic = False
                elif child.type == "code_inline":
                    run = paragraph.add_run(child.content)
                    run.bold = bold
                    run.italic = italic
                    _set_run_font(run, is_code=True)
                elif child.type == "text":
                    if child.content:
                        run = paragraph.add_run(child.content)
                        run.bold = bold
                        run.italic = italic
                        _set_run_font(run)
                elif child.type == "softbreak":
                    paragraph.add_run("\n")
                elif child.type == "hardbreak":
                    paragraph.add_run("\n")
                # 链接文字：只取文字，忽略链接
                elif child.type == "link_open":
                    pass
                elif child.type == "link_close":
                    pass

        # ── 解析 markdown ──
        md_parser = MarkdownIt("commonmark").enable("table")
        tokens = md_parser.parse(md_text)

        # 用栈跟踪列表嵌套
        list_depth = 0        # 当前列表嵌套深度
        list_type_stack = []  # 'bullet' | 'ordered'
        in_blockquote = False

        i = 0
        while i < len(tokens):
            t = tokens[i]

            # ── 标题 ──
            if t.type == "heading_open":
                level = int(t.tag[1])  # h1 → 1, h2 → 2, ...
                # 下一个 token 是 inline，包含标题文字
                if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                    h = doc.add_heading(level=level)
                    inline_tok = tokens[i + 1]
                    if inline_tok.children:
                        _add_inline_children(h, inline_tok.children)
                    else:
                        run = h.add_run(inline_tok.content)
                        _set_run_font(run)
                    i += 3  # heading_open, inline, heading_close
                    continue

            # ── 列表开始/结束 ──
            elif t.type == "bullet_list_open":
                list_depth += 1
                list_type_stack.append("bullet")
            elif t.type == "bullet_list_close":
                list_depth -= 1
                if list_type_stack:
                    list_type_stack.pop()
            elif t.type == "ordered_list_open":
                list_depth += 1
                list_type_stack.append("ordered")
            elif t.type == "ordered_list_close":
                list_depth -= 1
                if list_type_stack:
                    list_type_stack.pop()

            # ── 引用块 ──
            elif t.type == "blockquote_open":
                in_blockquote = True
            elif t.type == "blockquote_close":
                in_blockquote = False

            # ── 代码块 ──
            elif t.type == "fence":
                p = doc.add_paragraph()
                run = p.add_run(t.content.rstrip("\n"))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "Source Han Serif CN")

            # ── 段落（含列表项内容） ──
            elif t.type == "paragraph_open":
                # 找到对应的 inline token
                if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                    inline_tok = tokens[i + 1]

                    if list_depth > 0:
                        # 在列表中：根据类型和深度选择样式
                        current_type = list_type_stack[-1] if list_type_stack else "bullet"
                        if current_type == "ordered":
                            style_name = "List Number"
                        else:
                            style_name = "List Bullet"

                        # 深度 > 1 时尝试用 Word 的 "List Bullet 2" / "List Number 2" 等样式
                        if list_depth >= 2:
                            deeper_style = f"{style_name} {list_depth}"
                            if deeper_style in [s.name for s in doc.styles]:
                                style_name = deeper_style

                        p = doc.add_paragraph(style=style_name)

                        # 额外用缩进处理更深层级
                        if list_depth >= 2:
                            p.paragraph_format.left_indent = Cm(1.27 * (list_depth - 1))

                    elif in_blockquote:
                        p = doc.add_paragraph()
                        try:
                            p.style = doc.styles["Quote"]
                        except KeyError:
                            pass
                        p.paragraph_format.left_indent = Cm(1.27)
                    else:
                        p = doc.add_paragraph()

                    # 渲染内联内容
                    if inline_tok.children:
                        _add_inline_children(p, inline_tok.children)
                    elif inline_tok.content:
                        run = p.add_run(inline_tok.content)
                        _set_run_font(run)

                    i += 3  # paragraph_open, inline, paragraph_close
                    continue

            # ── 水平线 ──
            elif t.type == "hr":
                p = doc.add_paragraph()
                p.add_run("─" * 50)

            i += 1

        doc.save(str(out_path))
        return True
    except Exception:
        import traceback
        traceback.print_exc()
        return False


@router.post("/export_file")
def export_file(data: ExportFileRequest):
    """将文件保存到用户配置的输出目录（支持文本、二进制、格式转换）"""
    from app.utils.app_settings import get_note_output_dir
    import base64

    out_dir = get_note_output_dir()
    # 确保文件名安全
    safe_name = data.filename.strip().replace("/", "_").replace("\\", "_")
    if not safe_name:
        return R.error(msg="文件名不能为空")

    # 格式转换处理
    fmt = data.format.lower() if data.format else ""

    if fmt == "pdf":
        # 替换扩展名
        base_name = os.path.splitext(safe_name)[0]
        safe_name = f"{base_name}.pdf"
        file_path = out_dir / safe_name
        if _convert_md_to_pdf(data.content, file_path):
            return R.success(data={"path": str(file_path)}, msg="PDF 导出成功")
        else:
            return R.error(msg="PDF 转换失败，请检查后端控制台日志")

    if fmt == "docx":
        base_name = os.path.splitext(safe_name)[0]
        safe_name = f"{base_name}.docx"
        file_path = out_dir / safe_name
        if _convert_md_to_docx(data.content, file_path):
            return R.success(data={"path": str(file_path)}, msg="Word 导出成功")
        else:
            return R.error(msg="Word 转换失败，请安装 python-docx：pip install python-docx")

    # 默认：直接保存（md / html / svg 等文本，或 base64 二进制）
    file_path = out_dir / safe_name
    try:
        if data.is_base64:
            raw = base64.b64decode(data.content)
            file_path.write_bytes(raw)
        else:
            file_path.write_text(data.content, encoding="utf-8")
        return R.success(data={"path": str(file_path)}, msg="导出成功")
    except Exception as e:
        return R.error(msg=f"导出失败: {e}")